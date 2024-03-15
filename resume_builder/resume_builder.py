from docx import Document
from openai import AzureOpenAI

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)

# Example usage
docx_path = 'resume.docx'  # Replace with the path to your Word document
extracted_text = extract_text_from_docx(docx_path)

client = AzureOpenAI(
    azure_endpoint = 'https://dg1-sn.openai.azure.com/',
    api_key = 'e135ffd6b14f4957a3b94ac1c8ba91d4',
    api_version = '2024-02-01'
)

deployment_name = 'DG1-SN-GPT-35-TURBO'

job_title = input('What job do you need to optimize for? :: ')
# job_docx_path = f'{job_title}.docx'
# jobDesc = extracted_text = extract_text_from_docx(job_docx_path)

prompt = [
    #{'role':'system', 'content':'You are an AI job assistant for entry level positions that tazkes an old resume and a job description and creates a new resume with the same details as the old resume, but the new resume also contains key words from the job description.'},
    {'role':'system', 'content':'You are an AI assistant that takes a job title and resume text and returns new resume text that is the input resume text tailored to the job title'},
    {'role':'user', 'content':f'Old resume: {extracted_text}; Job Title: entry level {job_title}'}
]

response = client.chat.completions.create(model=deployment_name, messages=prompt)

new_resume_text = response.choices[0].message.content

def create_word_document(text, docx_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(docx_path)


new_docx_path = f'resume_{job_title}.docx'  # Path for the new Word document
create_word_document(new_resume_text, new_docx_path)

print(f'new resume for entry level {job_title} has been created.')